import os
import re
import hashlib
import tempfile
import json
import concurrent.futures
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import pickle
import matplotlib.pyplot as plt
import networkx as nx
import streamlit as st
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Document processing
import fitz  # PyMuPDF
import PyPDF2
from docx import Document

# Language model and embeddings
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

# Configuration constants
MAX_CHUNK_SIZE = 1000
OVERLAP_SIZE = 200
BATCH_SIZE = 10  # Number of pages to process at once for memory management

@st.cache_resource
def get_embeddings():
    # Load HuggingFace embeddings
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2"
    )

# Document Processing Functions
def compute_file_hash(file_bytes):
    """Compute MD5 hash for a file to use as cache key"""
    return hashlib.md5(file_bytes).hexdigest()

def extract_text_from_pdf_pymupdf(file_bytes, start_page=0, end_page=None):
    """Extract text from PDF using PyMuPDF (primary method)"""
    text_by_page = []
    
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
        total_pages = pdf_document.page_count
        end_page = end_page if end_page is not None else total_pages
        
        for page_num in range(start_page, min(end_page, total_pages)):
            page = pdf_document.load_page(page_num)
            text = page.get_text()
            text_by_page.append(text)
            
    return text_by_page

def extract_text_from_pdf_pypdf2(file_bytes, start_page=0, end_page=None):
    """Extract text from PDF using PyPDF2 (fallback method)"""
    text_by_page = []
    
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(file_bytes)
        temp_file_path = temp_file.name
    
    try:
        with open(temp_file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            end_page = end_page if end_page is not None else total_pages
            
            for page_num in range(start_page, min(end_page, total_pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_by_page.append(text)
    finally:
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
            
    return text_by_page

def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file"""
    text_by_page = []
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_file.write(file_bytes)
        temp_file_path = temp_file.name
    
    try:
        doc = Document(temp_file_path)
        full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Split into pseudo-pages (approx. 3000 chars per page)
        chars_per_page = 3000
        text_by_page = [full_text[i:i+chars_per_page] 
                       for i in range(0, len(full_text), chars_per_page)]
    finally:
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
            
    return text_by_page

def extract_text_from_txt(file_bytes):
    """Extract text from TXT file"""
    text = file_bytes.decode('utf-8')
    
    # Split into pseudo-pages (approx. 3000 chars per page)
    chars_per_page = 3000
    text_by_page = [text[i:i+chars_per_page] 
                   for i in range(0, len(text), chars_per_page)]
    
    return text_by_page

def process_document(uploaded_file, cache_dir, progress_bar=None):
    """Process document and create vector store"""
    # Read file bytes
    file_bytes = uploaded_file.getvalue()
    file_hash = compute_file_hash(file_bytes)
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    # Check if document is already cached
    cache_path = cache_dir / f"{file_hash}.pkl"
    if cache_path.exists():
        with open(cache_path, 'rb') as f:
            return pickle.load(f)
    
    # Extract text based on file type
    try:
        if file_extension == 'pdf':
            try:
                # Use PyMuPDF first
                text_by_page = []
                with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
                    total_pages = pdf_document.page_count
                    
                    # Process in batches to manage memory
                    for batch_start in range(0, total_pages, BATCH_SIZE):
                        batch_end = min(batch_start + BATCH_SIZE, total_pages)
                        
                        # Use parallel processing for the batch
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            future_to_page = {
                                executor.submit(
                                    lambda p: pdf_document.load_page(p).get_text(), 
                                    page_num
                                ): page_num 
                                for page_num in range(batch_start, batch_end)
                            }
                            
                            for future in concurrent.futures.as_completed(future_to_page):
                                page_num = future_to_page[future]
                                try:
                                    text = future.result()
                                    # Ensure text_by_page has enough slots
                                    while len(text_by_page) <= page_num - batch_start:
                                        text_by_page.append("")
                                    text_by_page[page_num - batch_start] = text
                                except Exception as exc:
                                    st.error(f"Error processing page {page_num}: {exc}")
                        
                        # Update progress
                        if progress_bar:
                            progress_bar.progress((batch_end) / total_pages)
            except Exception as e:
                st.warning(f"PyMuPDF extraction failed, falling back to PyPDF2: {e}")
                text_by_page = extract_text_from_pdf_pypdf2(file_bytes)
        elif file_extension == 'docx':
            text_by_page = extract_text_from_docx(file_bytes)
        elif file_extension == 'txt':
            text_by_page = extract_text_from_txt(file_bytes)
        else:
            st.error(f"Unsupported file format: {file_extension}")
            return None
    except Exception as e:
        st.error(f"Error processing document: {e}")
        return None
    
    # Create document object
    doc_info = {
        'filename': uploaded_file.name,
        'pages': len(text_by_page),
        'text_by_page': text_by_page,
        'file_hash': file_hash,
        'upload_time': st.session_state.get('current_time', ''),
    }
    
    # Cache the processed document
    with open(cache_path, 'wb') as f:
        pickle.dump(doc_info, f)
    
    return doc_info

def create_vector_store(document_info, cache_dir):
    """Create a vector store from document text"""
    
    # Check if we already have a vector store for this document
    vs_cache_path = cache_dir / f"{document_info['file_hash']}_vectorstore.pkl"
    if vs_cache_path.exists():
        with open(vs_cache_path, 'rb') as f:
            return pickle.load(f)
    
    # Combine all pages into one text
    all_text = '\n\n'.join(document_info['text_by_page'])
    
    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=MAX_CHUNK_SIZE,
        chunk_overlap=OVERLAP_SIZE,
        length_function=len
    )
    
    chunks = text_splitter.split_text(all_text)
    
    # Create documents for vectorstore
    documents = [
        LangchainDocument(page_content=chunk, metadata={"source": document_info['filename']})
        for chunk in chunks
    ]
    
    # Create vector store
    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(documents, embeddings)
    
    # Cache the vector store
    with open(vs_cache_path, 'wb') as f:
        pickle.dump(vectorstore, f)
    
    return vectorstore

def get_document_context(query, vectorstore, top_k=5):
    """Retrieve relevant context from the vector store based on query"""
    if vectorstore is None:
        return "No document loaded."
    
    results = vectorstore.similarity_search(query, k=top_k)
    context = "\n\n".join([doc.page_content for doc in results])
    return context

def generate_mind_map_data(mind_map_description):
    """Generate network graph data from mind map description"""
    # Extract central concept
    central_match = re.search(r"Central concept:\s*(.*?)(?:\n|$)", mind_map_description, re.IGNORECASE)
    central_node = central_match.group(1).strip() if central_match else "Main Topic"
    
    # Extract branches/nodes and connections
    nodes = [{"id": 0, "label": central_node, "group": 0}]
    edges = []
    
    # Look for patterns like "Branch: X" or "Connection: X - Y"
    branch_pattern = re.compile(r"Branch(?:\s\d+)?:\s*(.*?)(?:\n|$)", re.IGNORECASE)
    connection_pattern = re.compile(r"Connection(?:\s\d+)?:\s*(.*?)\s*-\s*(.*?)(?:\n|$)", re.IGNORECASE)
    
    branches = branch_pattern.findall(mind_map_description)
    for i, branch in enumerate(branches):
        branch = branch.strip()
        nodes.append({"id": i+1, "label": branch, "group": 1})
        edges.append({"from": 0, "to": i+1})
    
    connections = connection_pattern.findall(mind_map_description)
    for i, (source, target) in enumerate(connections):
        # Try to find existing nodes
        source = source.strip()
        target = target.strip()
        source_id = next((i for i, node in enumerate(nodes) if node["label"] == source), None)
        target_id = next((i for i, node in enumerate(nodes) if node["label"] == target), None)
        
        # If nodes don't exist, add them
        if source_id is None:
            source_id = len(nodes)
            nodes.append({"id": source_id, "label": source, "group": 2})
        if target_id is None:
            target_id = len(nodes)
            nodes.append({"id": target_id, "label": target, "group": 2})
        
        edges.append({"from": source_id, "to": target_id})
    
    return {"nodes": nodes, "edges": edges}

def parse_flashcards_from_text(text):
    """Parse flashcard data from text response"""
    try:
        # First try to see if it's valid JSON
        try:
            json_data = json.loads(text)
            if isinstance(json_data, list):
                return json_data
            elif isinstance(json_data, dict) and "flashcards" in json_data:
                return json_data["flashcards"]
        except json.JSONDecodeError:
            pass
        
        # If not valid JSON, try to parse structured text
        flashcards = []
        # Pattern for "Q: ... A: ..." format
        qa_pattern = re.compile(r"(?:Card\s*\d*:?\s*)?Q:\s*(.*?)\s*A:\s*(.*?)(?=(?:\n\s*(?:Card\s*\d*:?\s*)?Q:)|$)", re.DOTALL)
        matches = qa_pattern.findall(text)
        
        if matches:
            for q, a in matches:
                flashcards.append({"front": q.strip(), "back": a.strip()})
            return flashcards
        
        # Try another common format with numbered cards
        card_pattern = re.compile(r"(?:Card|Flashcard)\s*(\d+):\s*\n*Front:\s*(.*?)\s*\n*Back:\s*(.*?)(?=(?:\n\s*(?:Card|Flashcard)\s*\d+:)|$)", re.DOTALL)
        matches = card_pattern.findall(text)
        
        if matches:
            for _, front, back in matches:
                flashcards.append({"front": front.strip(), "back": back.strip()})
            return flashcards
            
        return []
    except Exception as e:
        st.error(f"Error parsing flashcards: {e}")
        return []

def parse_test_from_text(text):
    """Parse test data from text response with enhanced structure"""
    try:
        test_data = {
            "questions": [],
            "answer_key": {},
            "multiple_choice": [],
            "true_false": [],
            "short_answer": [],
            "essay": [],
            "instructions": "",
            "difficulty": "",
            "topic": ""
        }
        
        # Extract test metadata
        topic_match = re.search(r"Practice Test:\s*(.*?)(?=\n|\*\*)", text)
        if topic_match:
            test_data["topic"] = topic_match.group(1).strip()
        
        difficulty_match = re.search(r"\*\*Difficulty:\*\*\s*(.*?)(?=\n|\*\*)", text)
        if difficulty_match:
            test_data["difficulty"] = difficulty_match.group(1).strip()
        
        instructions_match = re.search(r"\*\*Instructions:\*\*\s*(.*?)(?=##|$)", text, re.DOTALL)
        if instructions_match:
            test_data["instructions"] = instructions_match.group(1).strip()
        
        # Parse Multiple Choice Questions
        mc_section = re.search(r"##\s*Multiple Choice Questions\s*(.*?)(?=##|$)", text, re.DOTALL)
        if mc_section:
            mc_text = mc_section.group(1)
            # Pattern for numbered questions with options
            mc_pattern = re.compile(r"(\d+)\.\s*(.*?)(?=\d+\.|$)", re.DOTALL)
            mc_questions = mc_pattern.findall(mc_text)
            
            for num, question_block in mc_questions:
                lines = question_block.strip().split('\n')
                question_text = lines[0].strip()
                options = []
                
                for line in lines[1:]:
                    line = line.strip()
                    if re.match(r'^[a-d]\)', line):
                        option = re.sub(r'^[a-d]\)\s*', '', line)
                        options.append(option)
                
                if question_text and options:
                    test_data["multiple_choice"].append({
                        "number": int(num),
                        "question": question_text,
                        "options": options
                    })
                    test_data["questions"].append(f"Q{num}: {question_text}")
        
        # Parse True/False Questions
        tf_section = re.search(r"##\s*True/False Questions\s*(.*?)(?=##|$)", text, re.DOTALL)
        if tf_section:
            tf_text = tf_section.group(1)
            tf_pattern = re.compile(r"(\d+)\.\s*(.*?)(?=\d+\.|$)", re.DOTALL)
            tf_questions = tf_pattern.findall(tf_text)
            
            for num, question_text in tf_questions:
                question_clean = question_text.strip().replace("(Explain your reasoning)", "").strip()
                test_data["true_false"].append({
                    "number": int(num),
                    "question": question_clean
                })
                test_data["questions"].append(f"Q{num}: {question_clean}")
        
        # Parse Short Answer Questions
        sa_section = re.search(r"##\s*Short Answer Questions\s*(.*?)(?=##|$)", text, re.DOTALL)
        if sa_section:
            sa_text = sa_section.group(1)
            sa_pattern = re.compile(r"(\d+)\.\s*(.*?)(?=\d+\.|$)", re.DOTALL)
            sa_questions = sa_pattern.findall(sa_text)
            
            for num, question_text in sa_questions:
                test_data["short_answer"].append({
                    "number": int(num),
                    "question": question_text.strip()
                })
                test_data["questions"].append(f"Q{num}: {question_text.strip()}")
        
        # Parse Essay Questions
        essay_section = re.search(r"##\s*Essay Questions\s*(.*?)(?=##|$)", text, re.DOTALL)
        if essay_section:
            essay_text = essay_section.group(1)
            essay_pattern = re.compile(r"(\d+)\.\s*(.*?)(?=\d+\.|$)", re.DOTALL)
            essay_questions = essay_pattern.findall(essay_text)
            
            for num, question_text in essay_questions:
                test_data["essay"].append({
                    "number": int(num),
                    "question": question_text.strip()
                })
                test_data["questions"].append(f"Q{num}: {question_text.strip()}")
        
        # Parse Answer Key
        answer_key_section = re.search(r"##\s*Answer Key\s*(.*?)$", text, re.DOTALL)
        if answer_key_section:
            answers_text = answer_key_section.group(1)
            
            # Parse different answer types
            answer_sections = {
                "Multiple Choice Answers": r"###\s*Multiple Choice Answers\s*(.*?)(?=###|$)",
                "True/False Answers": r"###\s*True/False Answers\s*(.*?)(?=###|$)",
                "Short Answer Answers": r"###\s*Short Answer Answers\s*(.*?)(?=###|$)",
                "Essay Question Answers": r"###\s*Essay Question Answers\s*(.*?)(?=###|$)"
            }
            
            for section_name, pattern in answer_sections.items():
                section_match = re.search(pattern, answers_text, re.DOTALL)
                if section_match:
                    section_text = section_match.group(1)
                    # Extract individual answers
                    answer_pattern = re.compile(r"(\d+)\.\s*\*\*(?:Correct\s+)?Answer:?\s*([^*]*?)\*\*\s*-\s*(.*?)(?=\d+\.|$)", re.DOTALL)
                    answers = answer_pattern.findall(section_text)
                    
                    for num, answer, explanation in answers:
                        test_data["answer_key"][num] = {
                            "answer": answer.strip(),
                            "explanation": explanation.strip()
                        }
        
        return test_data
        
    except Exception as e:
        st.error(f"Error parsing test: {e}")
        return {
            "questions": [],
            "answer_key": {},
            "multiple_choice": [],
            "true_false": [],
            "short_answer": [],
            "essay": [],
            "instructions": "",
            "difficulty": "",
            "topic": ""
        }

def parse_roadmap_from_text(text):
    """Parse roadmap data from text response"""
    try:
        roadmap_data = {
            "overview": "",
            "schedule": [],
            "milestones": [],
            "sections": []
        }
        
        # Extract overview
        overview_match = re.search(r"#+\s*Overview\s*(.*?)(?=#+|$)", text, re.DOTALL | re.IGNORECASE)
        if overview_match:
            roadmap_data["overview"] = overview_match.group(1).strip()
        
        # Extract day-by-day schedule
        schedule_pattern = re.compile(r"Day\s*(\d+):?\s*(.*?)(?=(?:Day\s*\d+)|$)", re.DOTALL | re.IGNORECASE)
        schedule_matches = schedule_pattern.findall(text)
        
        for day, content in schedule_matches:
            day_data = {
                "day": int(day),
                "content": content.strip(),
                "topics": [],
                "hours": 0
            }
            
            # Extract topics if available
            topics_match = re.search(r"Topics?:?\s*(.*?)(?=\n\n|$)", content, re.DOTALL)
            if topics_match:
                topics_text = topics_match.group(1)
                # Split by bullets or commas
                if "-" in topics_text:
                    day_data["topics"] = [t.strip() for t in topics_text.split("-") if t.strip()]
                else:
                    day_data["topics"] = [t.strip() for t in topics_text.split(",") if t.strip()]
            
            # Extract hours if available
            hours_match = re.search(r"(\d+(?:\.\d+)?)\s*hours?", content, re.IGNORECASE)
            if hours_match:
                day_data["hours"] = float(hours_match.group(1))
            
            roadmap_data["schedule"].append(day_data)
        
        # Extract milestones
        milestones_match = re.search(r"#+\s*Milestones\s*(.*?)(?=#+|$)", text, re.DOTALL | re.IGNORECASE)
        if milestones_match:
            milestones_text = milestones_match.group(1)
            # Try to extract individual milestones
            milestone_pattern = re.compile(r"(?:[\*\-•]\s*|\d+\.\s*)(.*?)(?=(?:[\*\-•]|\d+\.)|$)", re.DOTALL)
            milestones = milestone_pattern.findall(milestones_text)
            roadmap_data["milestones"] = [m.strip() for m in milestones if m.strip()]
        
        # Extract sections
        sections_match = re.search(r"#+\s*Sections\s*(.*?)(?=#+|$)", text, re.DOTALL | re.IGNORECASE)
        if sections_match:
            sections_text = sections_match.group(1)
            # Try to extract individual sections
            section_pattern = re.compile(r"(?:[\*\-•]\s*|\d+\.\s*)(.*?)(?=(?:[\*\-•]|\d+\.)|$)", re.DOTALL)
            sections = section_pattern.findall(sections_text)
            roadmap_data["sections"] = [s.strip() for s in sections if s.strip()]
        
        return roadmap_data
    except Exception as e:
        st.error(f"Error parsing roadmap: {e}")
        return {
            "overview": "Error parsing roadmap data",
            "schedule": [],
            "milestones": [],
            "sections": []
        }

# DSA Interview Utilities
def parse_dsa_questions_from_text(text):
    """Parse DSA questions from text response"""
    try:
        # Try to parse as JSON
        try:
            import json
            questions = json.loads(text)
            if isinstance(questions, list):
                return questions
            elif isinstance(questions, dict) and "questions" in questions:
                return questions["questions"]
        except json.JSONDecodeError:
            pass
        
        # Use regex to parse structured text
        import re
        questions = []
        
        # Pattern for "Question X: Title"
        question_pattern = re.compile(r"Question\s*(\d+):\s*(.*?)(?=\n\n|\n(?:Difficulty|Description):|\Z)", re.DOTALL)
        description_pattern = re.compile(r"Description:\s*(.*?)(?=\n\n|\n(?:Input|Output|Example|Difficulty|Topics):|\Z)", re.DOTALL)
        difficulty_pattern = re.compile(r"Difficulty:\s*(.*?)(?=\n\n|\n(?:Description|Input|Output|Example|Topics|Company):|\Z)", re.DOTALL)
        topics_pattern = re.compile(r"Topics:\s*(.*?)(?=\n\n|\n(?:Description|Input|Output|Example|Difficulty|Company):|\Z)", re.DOTALL)
        company_pattern = re.compile(r"Company:\s*(.*?)(?=\n\n|\n(?:Description|Input|Output|Example|Difficulty|Topics):|\Z)", re.DOTALL)
        
        # Find all question blocks
        blocks = re.split(r"\n\s*Question\s*\d+:", text)
        if len(blocks) > 1:  # First block is likely introduction or empty
            for i, block in enumerate(blocks[1:], 1):
                block_text = f"Question {i}:{block}"
                
                question_match = question_pattern.search(block_text)
                description_match = description_pattern.search(block_text)
                difficulty_match = difficulty_pattern.search(block_text)
                topics_match = topics_pattern.search(block_text)
                company_match = company_pattern.search(block_text)
                
                question = {
                    "id": i,
                    "title": question_match.group(2).strip() if question_match else f"Question {i}",
                    "description": description_match.group(1).strip() if description_match else block_text.strip(),
                    "difficulty": difficulty_match.group(1).strip() if difficulty_match else "Medium",
                    "topics": [t.strip() for t in topics_match.group(1).split(',')] if topics_match else [],
                    "companies": [c.strip() for c in company_match.group(1).split(',')] if company_match else []
                }
                questions.append(question)
        
        return questions
    except Exception as e:
        st.error(f"Error parsing DSA questions: {e}")
        return []

def calculate_dsa_progress_metrics(user_progress):
    """Calculate DSA practice progress metrics"""
    try:
        # Initialize metrics
        metrics = {
            "total_completed": 0,
            "total_attempted": 0,
            "by_difficulty": {"Easy": 0, "Medium": 0, "Hard": 0},
            "by_topic": {},
            "by_company": {},
            "success_rate": 0,
            "average_attempts": 0,
            "trending_topics": []
        }
        
        if not user_progress or not user_progress.get("questions"):
            return metrics
        
        questions = user_progress.get("questions", [])
        
        # Process each question
        for q in questions:
            if q.get("status") == "completed":
                metrics["total_completed"] += 1
            
            if q.get("attempts", 0) > 0:
                metrics["total_attempted"] += 1
            
            # Track by difficulty
            difficulty = q.get("difficulty", "Medium")
            if difficulty in metrics["by_difficulty"]:
                if q.get("status") == "completed":
                    metrics["by_difficulty"][difficulty] += 1
            
            # Track by topic
            for topic in q.get("topics", []):
                if topic not in metrics["by_topic"]:
                    metrics["by_topic"][topic] = {"completed": 0, "total": 0}
                
                metrics["by_topic"][topic]["total"] += 1
                if q.get("status") == "completed":
                    metrics["by_topic"][topic]["completed"] += 1
            
            # Track by company
            for company in q.get("companies", []):
                if company not in metrics["by_company"]:
                    metrics["by_company"][company] = {"completed": 0, "total": 0}
                
                metrics["by_company"][company]["total"] += 1
                if q.get("status") == "completed":
                    metrics["by_company"][company]["completed"] += 1
        
        # Calculate derived metrics
        if metrics["total_attempted"] > 0:
            metrics["success_rate"] = metrics["total_completed"] / metrics["total_attempted"] * 100
        
        # Calculate trending topics (topics with most questions)
        if metrics["by_topic"]:
            trending = sorted(metrics["by_topic"].items(), 
                             key=lambda x: x[1]["total"], 
                             reverse=True)[:5]
            metrics["trending_topics"] = [t[0] for t in trending]
        
        return metrics
    except Exception as e:
        st.error(f"Error calculating progress metrics: {e}")
        return {}

def parse_code_analysis(text):
    """Parse code analysis output from debugging agent"""
    try:
        import re
        
        analysis = {
            "bugs": [],
            "optimizations": [],
            "time_complexity": "",
            "space_complexity": "",
            "improved_code": ""
        }
        
        # Extract bugs section
        bugs_match = re.search(r"(?:Bugs|Issues)(?:\s+Found)?:\s*(.*?)(?=(?:\n\n|\n#|$))", text, re.DOTALL | re.IGNORECASE)
        if bugs_match:
            bugs_text = bugs_match.group(1)
            # Extract individual bugs (assuming they're numbered or bulleted)
            bug_items = re.findall(r"(?:^|\n)(?:\d+\.|\*|-)\s*(.*?)(?=(?:\n(?:\d+\.|\*|-|$))|\Z)", bugs_text, re.DOTALL)
            analysis["bugs"] = [bug.strip() for bug in bug_items if bug.strip()]
        
        # Extract optimizations section
        opt_match = re.search(r"(?:Optimizations|Improvements):\s*(.*?)(?=(?:\n\n|\n#|$))", text, re.DOTALL | re.IGNORECASE)
        if opt_match:
            opt_text = opt_match.group(1)
            # Extract individual optimizations
            opt_items = re.findall(r"(?:^|\n)(?:\d+\.|\*|-)\s*(.*?)(?=(?:\n(?:\d+\.|\*|-|$))|\Z)", opt_text, re.DOTALL)
            analysis["optimizations"] = [opt.strip() for opt in opt_items if opt.strip()]
        
        # Extract complexity analysis
        time_match = re.search(r"Time Complexity:\s*(.*?)(?=\n|$)", text, re.IGNORECASE)
        if time_match:
            analysis["time_complexity"] = time_match.group(1).strip()
        
        space_match = re.search(r"Space Complexity:\s*(.*?)(?=\n|$)", text, re.IGNORECASE)
        if space_match:
            analysis["space_complexity"] = space_match.group(1).strip()
        
        # Extract improved code section (assuming it's between code blocks or at the end)
        code_match = re.search(r"(?:Improved|Optimized) Code:?\s*(?:```[\w]*\n)?(.*?)(?:```|\Z)", text, re.DOTALL | re.IGNORECASE)
        if code_match:
            analysis["improved_code"] = code_match.group(1).strip()
        
        return analysis
    except Exception as e:
        st.error(f"Error parsing code analysis: {e}")
        return {}

def get_sample_dsa_questions():
    """Return a comprehensive list of DSA questions from various platforms for interview preparation"""
    return [
        {
            "id": 1,
            "title": "Two Sum",
            "description": "Given an array of integers nums and an integer target, return indices of the two numbers such that they add up to target.",
            "difficulty": "Easy",
            "topics": ["Array", "Hash Table"],
            "companies": ["Amazon", "Google", "Microsoft", "Facebook", "Adobe"],
            "link": "https://leetcode.com/problems/two-sum/",
            "platform": "LeetCode"
        },
        {
            "id": 2,
            "title": "Reverse Linked List",
            "description": "Given the head of a singly linked list, reverse the list, and return the reversed list.",
            "difficulty": "Easy",
            "topics": ["Linked List", "Recursion"],
            "companies": ["Amazon", "Apple", "Microsoft", "Facebook"],
            "link": "https://leetcode.com/problems/reverse-linked-list/",
            "platform": "LeetCode"
        },
        {
            "id": 3,
            "title": "Valid Parentheses",
            "description": "Given a string s containing just the characters '(', ')', '{', '}', '[' and ']', determine if the input string is valid.",
            "difficulty": "Easy",
            "topics": ["Stack", "String"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook"],
            "link": "https://leetcode.com/problems/valid-parentheses/",
            "platform": "LeetCode"
        },
        {
            "id": 4,
            "title": "Maximum Subarray",
            "description": "Given an integer array nums, find the subarray with the largest sum, and return its sum.",
            "difficulty": "Medium",
            "topics": ["Array", "Dynamic Programming", "Divide and Conquer"],
            "companies": ["Amazon", "Microsoft", "Apple", "Facebook", "Bloomberg"],
            "link": "https://leetcode.com/problems/maximum-subarray/",
            "platform": "LeetCode"
        },
        {
            "id": 5,
            "title": "Merge Intervals",
            "description": "Given an array of intervals where intervals[i] = [starti, endi], merge all overlapping intervals.",
            "difficulty": "Medium",
            "topics": ["Array", "Sorting"],
            "companies": ["Facebook", "Amazon", "Google", "Microsoft"],
            "link": "https://leetcode.com/problems/merge-intervals/",
            "platform": "LeetCode"
        },
        {
            "id": 6,
            "title": "3Sum",
            "description": "Given an integer array nums, return all the triplets [nums[i], nums[j], nums[k]] such that i != j, i != k, and j != k, and nums[i] + nums[j] + nums[k] == 0.",
            "difficulty": "Medium",
            "topics": ["Array", "Two Pointers", "Sorting"],
            "companies": ["Amazon", "Facebook", "Google", "Microsoft", "Apple", "Adobe"],
            "link": "https://leetcode.com/problems/3sum/",
            "platform": "LeetCode"
        },
        {
            "id": 7,
            "title": "LRU Cache",
            "description": "Design a data structure that follows the constraints of a Least Recently Used (LRU) cache.",
            "difficulty": "Medium",
            "topics": ["Hash Table", "Linked List", "Design"],
            "companies": ["Amazon", "Microsoft", "Facebook", "Google", "Uber"],
            "link": "https://leetcode.com/problems/lru-cache/",
            "platform": "LeetCode"
        },
        {
            "id": 8,
            "title": "Trapping Rain Water",
            "description": "Given n non-negative integers representing an elevation map where the width of each bar is 1, compute how much water it can trap after raining.",
            "difficulty": "Hard",
            "topics": ["Array", "Two Pointers", "Dynamic Programming", "Stack"],
            "companies": ["Amazon", "Google", "Apple", "Facebook", "Microsoft"],
            "link": "https://leetcode.com/problems/trapping-rain-water/",
            "platform": "LeetCode"
        },
        {
            "id": 9,
            "title": "Binary Tree Level Order Traversal",
            "description": "Given the root of a binary tree, return the level order traversal of its nodes' values.",
            "difficulty": "Medium",
            "topics": ["Tree", "Binary Tree", "BFS"],
            "companies": ["Amazon", "Microsoft", "Facebook", "Google"],
            "link": "https://leetcode.com/problems/binary-tree-level-order-traversal/",
            "platform": "LeetCode"
        },
        {
            "id": 10,
            "title": "Word Break",
            "description": "Given a string s and a dictionary of strings wordDict, return true if s can be segmented into a space-separated sequence of one or more dictionary words.",
            "difficulty": "Medium",
            "topics": ["Dynamic Programming", "Trie", "String"],
            "companies": ["Amazon", "Google", "Facebook", "Microsoft", "Uber"],
            "link": "https://leetcode.com/problems/word-break/",
            "platform": "LeetCode"
        },
        {
            "id": 11,
            "title": "Merge K Sorted Lists",
            "description": "You are given an array of k linked-lists lists, each linked-list is sorted in ascending order. Merge all the linked-lists into one sorted linked-list.",
            "difficulty": "Hard",
            "topics": ["Linked List", "Divide and Conquer", "Heap"],
            "companies": ["Amazon", "Google", "Microsoft", "Facebook"],
            "link": "https://leetcode.com/problems/merge-k-sorted-lists/",
            "platform": "LeetCode"
        },
        {
            "id": 12,
            "title": "Course Schedule",
            "description": "There are a total of numCourses courses you have to take, labeled from 0 to numCourses - 1. You are given an array prerequisites where prerequisites[i] = [ai, bi] indicates that you must take course bi first if you want to take course ai.",
            "difficulty": "Medium",
            "topics": ["Graph", "DFS", "BFS", "Topological Sort"],
            "companies": ["Google", "Amazon", "Facebook", "Microsoft"],
            "link": "https://leetcode.com/problems/course-schedule/",
            "platform": "LeetCode"
        },
        {
            "id": 13,
            "title": "Find Median from Data Stream",
            "description": "The median is the middle value in an ordered integer list. Design a data structure that supports adding new integers and finding the median.",
            "difficulty": "Hard",
            "topics": ["Heap", "Design", "Data Stream"],
            "companies": ["Amazon", "Google", "Microsoft", "Facebook"],
            "link": "https://leetcode.com/problems/find-median-from-data-stream/",
            "platform": "LeetCode"
        },
        {
            "id": 14,
            "title": "Min Stack",
            "description": "Design a stack that supports push, pop, top, and retrieving the minimum element in constant time.",
            "difficulty": "Easy",
            "topics": ["Stack", "Design"],
            "companies": ["Amazon", "Microsoft", "Google", "Bloomberg"],
            "link": "https://leetcode.com/problems/min-stack/",
            "platform": "LeetCode"
        },
        {
            "id": 15,
            "title": "Implement Trie (Prefix Tree)",
            "description": "A trie (pronounced as \"try\") or prefix tree is a tree data structure used to efficiently store and retrieve keys in a dataset of strings.",
            "difficulty": "Medium",
            "topics": ["Trie", "Design", "Tree"],
            "companies": ["Google", "Amazon", "Microsoft", "Facebook"],
            "link": "https://leetcode.com/problems/implement-trie-prefix-tree/",
            "platform": "LeetCode"
        },
        {
            "id": 16,
            "title": "Add Two Numbers",
            "description": "You are given two non-empty linked lists representing two non-negative integers. The digits are stored in reverse order, and each of their nodes contains a single digit.",
            "difficulty": "Medium",
            "topics": ["Linked List", "Math", "Recursion"],
            "companies": ["Amazon", "Microsoft", "Apple", "Facebook", "Adobe"],
            "link": "https://leetcode.com/problems/add-two-numbers/",
            "platform": "LeetCode"
        },
        {
            "id": 17,
            "title": "Median of Two Sorted Arrays",
            "description": "Given two sorted arrays nums1 and nums2 of size m and n respectively, return the median of the two sorted arrays.",
            "difficulty": "Hard",
            "topics": ["Array", "Binary Search", "Divide and Conquer"],
            "companies": ["Google", "Amazon", "Microsoft", "Facebook"],
            "link": "https://leetcode.com/problems/median-of-two-sorted-arrays/",
            "platform": "LeetCode"
        },
        {
            "id": 18,
            "title": "Best Time to Buy and Sell Stock",
            "description": "You are given an array prices where prices[i] is the price of a given stock on the ith day. You want to maximize your profit.",
            "difficulty": "Easy",
            "topics": ["Array", "Dynamic Programming"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook", "Apple"],
            "link": "https://leetcode.com/problems/best-time-to-buy-and-sell-stock/",
            "platform": "LeetCode"
        },
        {
            "id": 19,
            "title": "Kth Largest Element in an Array",
            "description": "Given an integer array nums and an integer k, return the kth largest element in the array.",
            "difficulty": "Medium",
            "topics": ["Array", "Divide and Conquer", "Sorting", "Heap"],
            "companies": ["Amazon", "Facebook", "Microsoft", "Google"],
            "link": "https://leetcode.com/problems/kth-largest-element-in-an-array/",
            "platform": "LeetCode"
        },
        {
            "id": 20,
            "title": "Permutations",
            "description": "Given an array nums of distinct integers, return all the possible permutations.",
            "difficulty": "Medium",
            "topics": ["Array", "Backtracking"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook"],
            "link": "https://leetcode.com/problems/permutations/",
            "platform": "LeetCode"
        },
        {
            "id": 21,
            "title": "Rotate Array",
            "description": "Given an integer array nums, rotate the array to the right by k steps, where k is non-negative.",
            "difficulty": "Easy",
            "topics": ["Array", "Math"],
            "companies": ["Amazon", "Microsoft", "Apple", "Facebook"],
            "link": "https://leetcode.com/problems/rotate-array/",
            "platform": "LeetCode"
        },
        {
            "id": 22,
            "title": "Sort Colors",
            "description": "Given an array nums with n objects colored red, white, or blue, sort them in-place so that objects of the same color are adjacent.",
            "difficulty": "Medium",
            "topics": ["Array", "Two Pointers", "Sorting"],
            "companies": ["Amazon", "Microsoft", "Apple", "Google"],
            "link": "https://leetcode.com/problems/sort-colors/",
            "platform": "LeetCode"
        },
        {
            "id": 23,
            "title": "Validate Binary Search Tree",
            "description": "Given the root of a binary tree, determine if it is a valid binary search tree (BST).",
            "difficulty": "Medium",
            "topics": ["Tree", "Binary Search Tree", "DFS"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook"],
            "link": "https://leetcode.com/problems/validate-binary-search-tree/",
            "platform": "LeetCode"
        },
        {
            "id": 24,
            "title": "Longest Palindromic Substring",
            "description": "Given a string s, return the longest palindromic substring in s.",
            "difficulty": "Medium",
            "topics": ["String", "Dynamic Programming"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook"],
            "link": "https://leetcode.com/problems/longest-palindromic-substring/",
            "platform": "LeetCode"
        },
        {
            "id": 25,
            "title": "Single Number",
            "description": "Given a non-empty array of integers nums, every element appears twice except for one. Find that single one.",
            "difficulty": "Easy",
            "topics": ["Array", "Bit Manipulation"],
            "companies": ["Amazon", "Microsoft", "Google", "Facebook"],
            "link": "https://leetcode.com/problems/single-number/",
            "platform": "LeetCode"
        },
        {
            "id": 26,
            "title": "Detect a Loop in Linked List",
            "description": "Given a linked list, check if the linked list has a loop or not.",
            "difficulty": "Easy",
            "topics": ["Linked List", "Two Pointers"],
            "companies": ["Amazon", "Microsoft", "Google", "Adobe"],
            "link": "https://www.geeksforgeeks.org/detect-loop-in-a-linked-list/",
            "platform": "GeeksforGeeks"
        },
        {
            "id": 27,
            "title": "K Closest Points to Origin",
            "description": "Given an array of points where points[i] = [xi, yi] represents a point on the X-Y plane and an integer k, return the k closest points to the origin (0, 0).",
            "difficulty": "Medium",
            "topics": ["Array", "Math", "Divide and Conquer", "Heap", "Sorting"],
            "companies": ["Amazon", "Facebook", "Google", "Microsoft"],
            "link": "https://leetcode.com/problems/k-closest-points-to-origin/",
            "platform": "LeetCode"
        },
        {
            "id": 28,
            "title": "Letter Combinations of a Phone Number",
            "description": "Given a string containing digits from 2-9 inclusive, return all possible letter combinations that the number could represent.",
            "difficulty": "Medium",
            "topics": ["Hash Table", "String", "Backtracking"],
            "companies": ["Amazon", "Google", "Facebook", "Microsoft"],
            "link": "https://leetcode.com/problems/letter-combinations-of-a-phone-number/",
            "platform": "LeetCode"
        },
        {
            "id": 29,
            "title": "Minimum Window Substring",
            "description": "Given two strings s and t, return the minimum window in s which will contain all the characters in t.",
            "difficulty": "Hard",
            "topics": ["Hash Table", "String", "Sliding Window"],
            "companies": ["Amazon", "Google", "Facebook", "Microsoft"],
            "link": "https://leetcode.com/problems/minimum-window-substring/",
            "platform": "LeetCode"
        },
        {
            "id": 30,
            "title": "Meeting Rooms II",
            "description": "Given an array of meeting time intervals, find the minimum number of conference rooms required.",
            "difficulty": "Medium",
            "topics": ["Sorting", "Greedy", "Heap"],
            "companies": ["Amazon", "Google", "Facebook", "Microsoft", "Uber"],
            "link": "https://leetcode.com/problems/meeting-rooms-ii/",
            "platform": "LeetCode Premium"
        }
    ]

# Safety and Input Validation Functions
def validate_user_input(user_input: str) -> Dict[str, Any]:
    """
    Validate user input for safety and appropriateness
    Returns a dictionary with validation results
    """
    validation_result = {
        "is_valid": True,
        "warnings": [],
        "filtered_input": user_input,
        "risk_level": "low"  # low, medium, high
    }
    
    # Check for empty or too short input
    if not user_input or len(user_input.strip()) < 3:
        validation_result["is_valid"] = False
        validation_result["warnings"].append("Please provide a more detailed question.")
        return validation_result
    
    # Check for extremely long input (potential abuse)
    if len(user_input) > 2000:
        validation_result["is_valid"] = False
        validation_result["warnings"].append("Question is too long. Please keep it under 2000 characters.")
        return validation_result
    
    # Check for inappropriate content patterns
    inappropriate_patterns = [
        r'\b(hack|crack|piracy|illegal)\b',
        r'\b(personal\s+info|private\s+data|password)\b',
        r'\b(violence|weapon|bomb|kill)\b',
        r'\b(drug|substance\s+abuse)\b'
    ]
    
    for pattern in inappropriate_patterns:
        if re.search(pattern, user_input.lower()):
            validation_result["is_valid"] = False
            validation_result["risk_level"] = "high"
            validation_result["warnings"].append("This question contains inappropriate content. Please focus on educational topics.")
            return validation_result
    
    # Check for potential prompt injection attempts
    injection_patterns = [
        r'ignore\s+previous\s+instructions',
        r'system\s+prompt',
        r'you\s+are\s+now',
        r'forget\s+everything',
        r'new\s+role',
        r'act\s+as'
    ]
    
    for pattern in injection_patterns:
        if re.search(pattern, user_input.lower()):
            validation_result["risk_level"] = "medium"
            validation_result["warnings"].append("Please ask educational questions related to your study material.")
            validation_result["filtered_input"] = re.sub(pattern, "[filtered]", user_input.lower())
    
    return validation_result

def is_educational_question(question: str) -> bool:
    """
    Determine if a question is educational in nature
    """
    educational_keywords = [
        'what', 'how', 'why', 'when', 'where', 'explain', 'define', 'describe',
        'compare', 'contrast', 'analyze', 'summarize', 'discuss', 'evaluate',
        'calculate', 'solve', 'prove', 'demonstrate', 'illustrate', 'example',
        'concept', 'theory', 'principle', 'formula', 'equation', 'method'
    ]
    
    question_lower = question.lower()
    return any(keyword in question_lower for keyword in educational_keywords)

def format_chatbot_response(response: str) -> str:
    """
    Format and enhance chatbot response for better presentation
    """
    if not response:
        return "I apologize, but I couldn't generate a proper response. Please try rephrasing your question."
    
    # Ensure response has proper structure
    if not response.strip().endswith(('.', '!', '?')):
        response += "."
    
    # Add helpful closing if not already present
    if not any(phrase in response.lower() for phrase in ['follow-up', 'next', 'further', 'more']):
        response += "\n\n💡 **Need clarification?** Feel free to ask follow-up questions about this topic!"
    
    return response

def sanitize_file_content(content: str) -> str:
    """
    Sanitize file content to remove potentially harmful or irrelevant information
    """
    # Remove potential personal information patterns
    # Email addresses
    content = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL_REMOVED]', content)
    
    # Phone numbers (basic pattern)
    content = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE_REMOVED]', content)
    
    # Credit card patterns (basic)
    content = re.sub(r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b', '[CARD_REMOVED]', content)
    
    return content
